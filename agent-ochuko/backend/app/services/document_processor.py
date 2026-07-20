"""
Document Processor Service for Agent Ochuko.
Handles PDF image extraction (signatures/letterheads), DOCX text replacement,
signatory block updates, and image placement using PyMuPDF (fitz), python-docx, and Pillow.
"""
import io
import os
import logging
from typing import List, Dict, Any, Optional, Tuple
import fitz  # PyMuPDF
import docx
from docx.shared import Inches, Pt
from PIL import Image

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Helper methods for PDF extraction and DOCX document editing."""

    @staticmethod
    def extract_images_from_pdf(pdf_path: str, output_dir: str) -> List[str]:
        """
        Extracts all embedded images from a PDF file using PyMuPDF (fitz).
        Returns list of extracted image file paths.
        """
        if not os.path.exists(pdf_path):
            logger.error(f"PDF file not found: {pdf_path}")
            return []

        os.makedirs(output_dir, exist_ok=True)
        extracted_paths = []

        try:
            doc = fitz.open(pdf_path)
            for page_index in range(len(doc)):
                page = doc[page_index]
                image_list = page.get_images(full=True)

                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    img_filename = f"extracted_p{page_index+1}_img{img_index+1}.{image_ext}"
                    img_path = os.path.join(output_dir, img_filename)

                    with open(img_path, "wb") as f:
                        f.write(image_bytes)

                    extracted_paths.append(img_path)
                    logger.info(f"Extracted PDF image: {img_path}")

            doc.close()
        except Exception as e:
            logger.error(f"Failed to extract images from PDF {pdf_path}: {e}")

        return extracted_paths

    @staticmethod
    def replace_signatory_in_docx(docx_path: str, output_docx_path: str, old_signatory: str, new_signatory: str, new_title: str) -> bool:
        """
        Replaces old signatory text in a DOCX document with new_signatory and new_title.
        """
        if not os.path.exists(docx_path):
            logger.error(f"DOCX file not found: {docx_path}")
            return False

        try:
            doc = docx.Document(docx_path)

            for paragraph in doc.paragraphs:
                if old_signatory.lower() in paragraph.text.lower():
                    # Replace text in runs
                    for run in paragraph.runs:
                        if old_signatory.lower() in run.text.lower():
                            run.text = run.text.replace(old_signatory, new_signatory)
                    
                    # Add title if missing
                    if new_title and new_title.lower() not in paragraph.text.lower():
                        paragraph.text = f"{new_signatory}\n{new_title}"

            doc.save(output_docx_path)
            logger.info(f"Updated signatory in DOCX: {output_docx_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to replace signatory in DOCX: {e}")
            return False

    @staticmethod
    def insert_signature_image(docx_path: str, output_docx_path: str, signature_img_path: str, target_signatory: str) -> bool:
        """
        Inserts a signature image directly above the signatory paragraph in a DOCX document.
        """
        if not os.path.exists(docx_path) or not os.path.exists(signature_img_path):
            logger.error(f"Missing input DOCX or signature image file.")
            return False

        try:
            doc = docx.Document(docx_path)
            inserted = False

            for i, paragraph in enumerate(doc.paragraphs):
                if target_signatory.lower() in paragraph.text.lower():
                    # Insert picture in paragraph right before signatory
                    p_sig = paragraph.insert_paragraph_before()
                    run = p_sig.add_run()
                    run.add_picture(signature_img_path, width=Inches(1.8))
                    inserted = True
                    break

            if not inserted and doc.paragraphs:
                # Append at bottom if target not found explicitly
                p_last = doc.add_paragraph()
                r = p_last.add_run()
                r.add_picture(signature_img_path, width=Inches(1.8))

            doc.save(output_docx_path)
            logger.info(f"Inserted signature image into DOCX: {output_docx_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to insert signature image: {e}")
            return False


document_processor = DocumentProcessor()
