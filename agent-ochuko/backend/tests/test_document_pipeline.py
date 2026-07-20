"""
Automated Integration Test Harness for Document Processing & Extraction Pipeline.
Simulates PDF signature extraction, letterhead processing, DOCX text replacement,
signatory block update to Mr. Ochuko Ederagoghene COO, and verification gate checks.
"""
import os
import pytest
import docx
import fitz  # PyMuPDF
from PIL import Image, ImageDraw
from app.services.document_processor import document_processor
from app.core.verification_gates import verification_gates


@pytest.fixture
def mock_document_environment(tmp_path):
    """Creates mock PDF employment letter with signature, mock PDF letterhead, and mock DOCX letter."""
    env_dir = tmp_path / "doc_test_env"
    env_dir.mkdir()

    # 1. Create a mock signature image using Pillow
    sig_img_path = env_dir / "signature_mock.png"
    img = Image.new("RGBA", (200, 80), color=(255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    draw.line((10, 50, 190, 20), fill="blue", width=3)
    draw.text((20, 55), "Ochuko Ederagoghene", fill="black")
    img.save(sig_img_path)

    # 2. Create PDF with embedded signature image (simulating employment letter)
    pdf_sig_path = env_dir / "EKO ORGANIC OFFER LETTER OF EMPLOYMENT - IT_SIGNED.pdf"
    doc = fitz.open()
    page = doc.new_page()
    rect = fitz.Rect(100, 100, 300, 180)
    page.insert_image(rect, filename=str(sig_img_path))
    doc.save(str(pdf_sig_path))
    doc.close()

    # 3. Create DOCX letter (simulating original letter with Engr. Felix signatory)
    docx_input_path = env_dir / "letter to be put in the letter head.docx"
    doc_x = docx.Document()
    doc_x.add_heading("EXPRESSION OF INTEREST FOR AMMONIA PLANT", level=1)
    doc_x.add_paragraph("We write to formally express our interest in the acquisition of a used Ammonia Production Plant...")
    doc_x.add_paragraph("Yours faithfully,\nFor EUROCLEAN MULTI CONCEPT LIMITED")
    doc_x.add_paragraph("Engr. Felix Ederagoghene")
    doc_x.save(str(docx_input_path))

    return {
        "dir": env_dir,
        "pdf_sig_path": str(pdf_sig_path),
        "docx_input_path": str(docx_input_path),
        "sig_img_path": str(sig_img_path)
    }


def test_full_document_processing_pipeline(mock_document_environment):
    env = mock_document_environment
    out_dir = env["dir"] / "extracted"

    # Step 1: Extract signature image from PDF via PyMuPDF
    extracted_images = document_processor.extract_images_from_pdf(env["pdf_sig_path"], str(out_dir))
    assert len(extracted_images) > 0
    extracted_sig = extracted_images[0]
    assert os.path.exists(extracted_sig)

    # Step 2: Replace signatory in DOCX (Engr. Felix -> Mr. Ochuko Ederagoghene, COO)
    docx_replaced = env["dir"] / "letter_replaced.docx"
    success_replace = document_processor.replace_signatory_in_docx(
        docx_path=env["docx_input_path"],
        output_docx_path=str(docx_replaced),
        old_signatory="Engr. Felix Ederagoghene",
        new_signatory="Mr. Ochuko Ederagoghene",
        new_title="Chief Operating Officer"
    )
    assert success_replace is True

    # Step 3: Insert extracted signature image into DOCX above signatory
    docx_final = env["dir"] / "letter_final_signed.docx"
    success_insert = document_processor.insert_signature_image(
        docx_path=str(docx_replaced),
        output_docx_path=str(docx_final),
        signature_img_path=extracted_sig,
        target_signatory="Mr. Ochuko Ederagoghene"
    )
    assert success_insert is True
    assert os.path.exists(docx_final)

    # Step 4: Verify output DOCX header magic bytes using verification_gates
    is_valid_hdr, err_hdr = verification_gates.verify_document_header(str(docx_final))
    assert is_valid_hdr is True
    assert err_hdr is None

    # Step 5: Read finalized DOCX content and confirm signatory text
    final_doc = docx.Document(str(docx_final))
    full_text = "\n".join(p.text for p in final_doc.paragraphs)
    assert "Mr. Ochuko Ederagoghene" in full_text
    assert "Chief Operating Officer" in full_text
    assert "Engr. Felix" not in full_text
